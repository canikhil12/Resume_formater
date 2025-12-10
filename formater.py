import os
import re
import json
import random
from io import BytesIO

import docx
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph
import streamlit as st
from openai import OpenAI

# ============================================================
#  OpenAI client  (dummy placeholder – replace with your real key)
# ============================================================

client = OpenAI(api_key="your_api_key")


# ============================================================
#  DOCX helpers
# ============================================================
def read_full_text(doc: docx.Document) -> str:
    """Return plain text of the whole document (for ATS)."""
    return "\n".join(p.text for p in doc.paragraphs)


def is_bullet_paragraph(para: docx.text.paragraph.Paragraph) -> bool:
    """
    Detect whether a paragraph is a bullet/numbered list item
    using the underlying XML numPr (so we don't rely on '•' characters).
    """
    pPr = para._p.pPr
    return pPr is not None and pPr.numPr is not None


def get_experience_bullet_paragraphs(doc: docx.Document):
    """
    Find all bullet paragraphs only inside the Professional Experience section.

    We:
    - Start after a heading that contains 'Professional Experience'
    - Stop when we hit another main section like Certifications, Projects,
      Technical Skills, Skills, Education, etc.
    - Return a list of Paragraph objects that are bullets.
    """
    bullets = []
    inside_exp = False

    section_starts = ["PROFESSIONAL EXPERIENCE"]
    section_ends = [
        "CERTIFICATIONS & ACHIEVEMENTS",
        "CERTIFICATIONS",
        "PROJECTS",
        "TECHNICAL SKILLS",
        "SKILLS",
        "EDUCATION",
    ]

    for para in doc.paragraphs:
        txt = para.text.strip()
        upper = txt.upper()

        # Start of experience section
        if any(start in upper for start in section_starts):
            inside_exp = True
            continue

        # Reached another big section -> stop collecting
        if inside_exp and any(upper.startswith(end) for end in section_ends):
            break

        if not inside_exp:
            continue

        # Inside Professional Experience: collect bullet paragraphs only
        if is_bullet_paragraph(para) and txt:
            bullets.append(para)

    return bullets


def replace_paragraph_text_preserving_format(
    para: docx.text.paragraph.Paragraph, new_text: str
):
    """
    Replace the visible text of a paragraph while **preserving**:
    - List/bullet formatting
    - Styles
    - Alignment
    """
    # Clear existing runs
    for run in para.runs:
        run.text = ""

    # Add a single run with new text
    para.add_run(new_text)


def insert_paragraph_after(paragraph: Paragraph, text: str = "") -> Paragraph:
    """
    Insert a new paragraph *after* the given paragraph and return it.
    """
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if text:
        new_para.add_run(text)
    return new_para


# ============================================================
#  Skills section updater (auto-categorize + Others:)
# ============================================================
def categorize_skill(skill: str) -> str:
    """
    Heuristic categorization of a single skill string into a category.
    Returns one of the category names or "Tools" as a default.
    """
    s = skill.lower()

    categories = {
        "Languages": [
            "python", "java", "c++", "c#", "c ", "sql", "scala", "r ",
            "javascript", "typescript", "js", "go", "rust"
        ],
        "Databases": [
            "postgres", "postgresql", "mysql", "mariadb", "mongodb",
            "redshift", "snowflake", "oracle", "sql server", "dynamodb",
            "bigquery"
        ],
        "Data Engineering": [
            "spark", "databricks", "airflow", "data factory", "synapse",
            "kafka", "etl", "elt", "pipeline"
        ],
        "Machine Learning": [
            "machine learning", "deep learning", "transformer", "transformers",
            "recommendation", "recommender", "ranking", "personalization",
            "llm", "llms", "foundation model", "xgboost", "lightgbm",
            "random forest", "regression", "clustering", "reinforcement learning"
        ],
        "MLOps": [
            "mlops", "mlflow", "kubernetes", "docker", "ci/cd", "ci-cd",
            "github actions", "kubeflow", "model monitoring", "feature store"
        ],
        "Visualization": [
            "tableau", "power bi", "looker", "superset", "data studio",
            "lookerstudio"
        ],
        "Cloud Platforms": [
            "aws", "azure", "gcp", "google cloud", "s3", "ec2", "lambda",
            "redshift", "snowflake", "databricks", "adls"
        ],
        "Tools": [
            "git", "excel", "jira", "confluence", "jenkins", "intellij",
            "pycharm", "vscode"
        ],
    }

    for cat, kw_list in categories.items():
        for kw in kw_list:
            if kw in s:
                return cat

    # Default bucket
    return "Tools"


def update_skills_section(
    doc: docx.Document,
    suggestions: dict,
    keyword_analysis: dict,
):
    """
    Update the TECHNICAL SKILLS / SKILLS section to include:
    - Auto-categorized existing skills into labeled lines
    - "Others:" line with suggested skills + missing JD keywords

    Formatting (no bullets, just lines):

    Technical Skills
    Languages: Python, SQL, Java
    Databases: PostgreSQL, MySQL
    ...
    Others: Transformers, Feature Stores, Counterfactual Explainability, Personalized Ranking
    """
    # Gather new skills to add under Others
    suggested_skills = suggestions.get("suggested_skills_to_add", []) or []
    core_missing = keyword_analysis.get("core_missing", []) or []
    secondary_missing = keyword_analysis.get("secondary_missing", []) or []

    others_raw = suggested_skills + core_missing + secondary_missing

    # Clean + dedupe Others
    others_clean = []
    seen_others = set()
    for s in others_raw:
        sk = s.strip()
        if not sk:
            continue
        key = sk.lower()
        if key in seen_others:
            continue
        seen_others.add(key)
        others_clean.append(sk)

    # Locate SKILLS section
    section_starts = ["TECHNICAL SKILLS", "SKILLS"]
    section_ends = [
        "PROFESSIONAL EXPERIENCE",
        "EXPERIENCE",
        "PROJECTS",
        "CERTIFICATIONS & ACHIEVEMENTS",
        "CERTIFICATIONS",
        "EDUCATION",
    ]

    inside_skills = False
    skills_paras = []

    for para in doc.paragraphs:
        txt = para.text.strip()
        upper = txt.upper()

        if any(start == upper for start in section_starts):
            inside_skills = True
            continue

        if inside_skills and any(upper.startswith(end) for end in section_ends):
            break

        if inside_skills:
            skills_paras.append(para)

    if not skills_paras:
        # No dedicated skills section, skip silently
        return

    # Combine existing skills text
    combined_existing = " ".join(p.text for p in skills_paras)

    if not combined_existing.strip() and not others_clean:
        return

    # Normalize & split existing skills
    norm_text = combined_existing
    norm_text = norm_text.replace("\n", " ")
    norm_text = norm_text.replace(" and ", ", ")
    norm_text = norm_text.replace("And ", ", ")
    norm_text = norm_text.replace(":", ", ")
    norm_text = norm_text.replace(";", ", ")

    existing_tokens = [
        s.strip()
        for s in re.split(r"[,\u2022\|/]", norm_text)
        if s.strip()
    ]

    # Categorize existing skills
    categories_order = [
        "Languages",
        "Databases",
        "Data Engineering",
        "Machine Learning",
        "MLOps",
        "Visualization",
        "Cloud Platforms",
        "Tools",
    ]

    categorized = {cat: [] for cat in categories_order}

    seen_existing = set()
    for sk in existing_tokens:
        key = sk.lower()
        if key in seen_existing:
            continue
        seen_existing.add(key)
        cat = categorize_skill(sk)
        if cat not in categorized:
            cat = "Tools"
        categorized[cat].append(sk)

    # Build line texts
    lines = []
    for cat in categories_order:
        skills_list = categorized[cat]
        if not skills_list:
            continue
        # preserve order as they appeared
        skills_str = ", ".join(skills_list)
        lines.append(f"{cat}: {skills_str}")

    # Append Others line if we have anything
    if others_clean:
        others_str = ", ".join(others_clean)
        # Exactly the format you requested: "Others:Transformers, Feature Stores, ..."
        # (keeping a space after colon so it still reads nicely)
        lines.append(f"Others: {others_str}")

    # Write lines into paragraphs
    if not lines:
        return

    # Ensure we have enough paragraphs – reuse existing, create extra if needed
    first_para = skills_paras[0]
    replace_paragraph_text_preserving_format(first_para, lines[0])

    current_para = first_para
    for i in range(1, len(lines)):
        if i < len(skills_paras):
            para = skills_paras[i]
        else:
            para = insert_paragraph_after(current_para)
            skills_paras.append(para)
        replace_paragraph_text_preserving_format(para, lines[i])
        current_para = para

    # Clear any extra old paras beyond the number of lines
    for extra_para in skills_paras[len(lines):]:
        for run in extra_para.runs:
            run.text = ""


# ============================================================
#  LLM helpers – JD keyword buckets & ATS
# ============================================================
def extract_jd_keyword_buckets(job_description: str) -> dict:
    """
    Ask the LLM to produce JSON buckets of keywords:
    {
      "core_keywords": [...],
      "secondary_keywords": [...],
      "nice_to_have": [...]
    }
    """
    prompt = f"""
You are an ATS / hiring-engine simulator.

From the following job description, extract 3 buckets of skills / keywords:

1. "core_keywords"      – must-have skills / tools / concepts (around 4–6 items)
2. "secondary_keywords" – important but not absolutely required (around 4–6 items)
3. "nice_to_have"       – differentiators / plus points (around 3–5 items)

Rules:
- Focus on hard skills, tools, technologies, and domain phrases.
- Exclude generic soft skills like "communication", "team player", "fast learner".
- Return a **valid JSON object only**, no extra text.

JOB DESCRIPTION:
{job_description}
"""

    response = client.chat.completions.create(
        model="gpt-4o",
        response_format={"type": "json_object"},
        messages=[{"role": "user", "content": prompt}],
        temperature=0.0,
    )
    raw = response.choices[0].message.content.strip()

    try:
        buckets = json.loads(raw)
    except Exception:
        flat = [s.strip() for s in re.split(r"[,\n]", raw) if s.strip()]
        buckets = {
            "core_keywords": flat,
            "secondary_keywords": [],
            "nice_to_have": [],
        }

    for k in ["core_keywords", "secondary_keywords", "nice_to_have"]:
        buckets.setdefault(k, [])
        buckets[k] = [s.strip() for s in buckets[k] if s.strip()]

    return buckets


def compute_enterprise_ats(resume_text: str, job_description: str):
    """
    Enterprise-style ATS:

    - Uses JD keyword buckets (core/secondary/nice_to_have) from LLM.
    - Counts presence + frequency in resume.
    - Computes weighted score.
    - Asks LLM to suggest additional bullets & skills based on missing keywords.
    """
    buckets = extract_jd_keyword_buckets(job_description)

    resume_lower = resume_text.lower()

    def count_matches(keyword: str) -> int:
        # Substring-based, case-insensitive match (more forgiving)
        kw = keyword.lower().strip()
        if not kw:
            return 0
        return resume_lower.count(kw)

    def analyze_bucket(keywords):
        present = {}
        missing = []
        for kw in keywords:
            if not kw:
                continue
            freq = count_matches(kw)
            if freq > 0:
                present[kw] = freq
            else:
                missing.append(kw)
        return present, missing

    core_present, core_missing = analyze_bucket(buckets["core_keywords"])
    sec_present, sec_missing = analyze_bucket(buckets["secondary_keywords"])
    nice_present, nice_missing = analyze_bucket(buckets["nice_to_have"])

    def coverage(present_dict, total_keywords, weight_presence=1.0, weight_freq=0.2):
        if total_keywords == 0:
            return 0.0
        presence_ratio = len(present_dict) / total_keywords
        total_freq = sum(min(3, f) for f in present_dict.values())
        max_freq = 3 * total_keywords
        freq_ratio = total_freq / max_freq if max_freq > 0 else 0
        return weight_presence * presence_ratio + weight_freq * freq_ratio

    core_cov = coverage(core_present, len(buckets["core_keywords"]))
    sec_cov = coverage(sec_present, len(buckets["secondary_keywords"]))
    nice_cov = coverage(nice_present, len(buckets["nice_to_have"]))

    score = (
        0.6 * core_cov +
        0.3 * sec_cov +
        0.1 * nice_cov
    ) * 100
    score = round(score, 1)

    # Suggestions (bullets + skills)
    suggestions_prompt = f"""
You are helping a candidate upgrade their resume for this job.

JOB DESCRIPTION:
{job_description}

RESUME (text):
{resume_text}

JD KEYWORD BUCKETS:
{json.dumps(buckets, indent=2)}

Missing or weak keywords (good to weave into experience bullets and skills):

CORE missing: {core_missing}
SECONDARY missing: {sec_missing}
NICE_TO_HAVE missing: {nice_missing}

Return a short JSON with:
{{
  "suggested_experience_bullets": [
    "...",  // 3-5 bullet sentences that the user could add or adapt
  ],
  "suggested_skills_to_add": [
    "...",  // 5-8 concrete skill phrases to add in Technical Skills / Others
  ]
}}

Rules:
- Keep every bullet one sentence, concise but impact-focused.
- Only use tools / domains that are plausible for a data / analytics profile.
- No intro text, just valid JSON.
"""

    sug_response = client.chat.completions.create(
        model="gpt-4o",
        response_format={"type": "json_object"},
        messages=[{"role": "user", "content": suggestions_prompt}],
        temperature=0.4,
    )
    sug_raw = sug_response.choices[0].message.content.strip()

    suggested_bullets = []
    suggested_skills = []

    try:
        sug_obj = json.loads(sug_raw)
        suggested_bullets = [
            b.strip() for b in sug_obj.get("suggested_experience_bullets", []) if b.strip()
        ]
        suggested_skills = [
            s.strip() for s in sug_obj.get("suggested_skills_to_add", []) if s.strip()
        ]
    except Exception:
        for line in sug_raw.splitlines():
            line = line.strip(" -*•")
            if not line:
                continue
            if len(suggested_bullets) < 5:
                suggested_bullets.append(line)
            else:
                suggested_skills.append(line)

    keyword_analysis = {
        "core_keywords": buckets["core_keywords"],
        "secondary_keywords": buckets["secondary_keywords"],
        "nice_to_have": buckets["nice_to_have"],
        "core_present": core_present,
        "core_missing": core_missing,
        "secondary_present": sec_present,
        "secondary_missing": sec_missing,
        "nice_present": nice_present,
        "nice_missing": nice_missing,
    }

    suggestions = {
        "suggested_experience_bullets": suggested_bullets,
        "suggested_skills_to_add": suggested_skills,
    }

    return score, keyword_analysis, suggestions


# ============================================================
#  Experience bullet rewriter (weaves in missing keywords)
# ============================================================
def rewrite_experience_bullets(
    original_bullets,
    job_description: str,
    keyword_analysis: dict,
):
    """
    Use OpenAI to rewrite ONLY the bullet contents in Professional Experience.

    - Keeps same number of bullets (helps keep resume ~1 page).
    - Tries to naturally weave in missing core/secondary keywords.
    """
    joined_bullets = "\n".join(original_bullets)
    bullet_count = len(original_bullets)

    core_missing = keyword_analysis["core_missing"]
    sec_missing = keyword_analysis["secondary_missing"]

    prompt = f"""
You are a senior resume optimization engine.

Rewrite the following **professional experience bullet points** so they better match
the job description and cover as many of the missing CORE and SECONDARY keywords
below as is honest and realistic.

CURRENT BULLETS:
{joined_bullets}

JOB DESCRIPTION:
{job_description}

JD MISSING KEYWORDS TO COVER IF POSSIBLE:
- CORE missing: {core_missing}
- SECONDARY missing: {sec_missing}

RULES:
- KEEP the same number of bullet points: exactly {bullet_count} lines.
- Each line must be a single bullet sentence.
- Do NOT add headings, job titles, company names, or dates.
- Use strong, varied action verbs.
- Include metrics/impact where reasonable.
- ONLY use tools and technologies that are plausible for a data / analytics profile.
- Weave in the above missing keywords **naturally** where it makes sense.
- RETURN:
  * Plain text
  * Exactly {bullet_count} lines
  * No numbering, no bullet characters (no "•", "-", "1.", etc.)

Rewritten bullet points:
"""

    response = client.chat.completions.create(
        model="gpt-4o",
        messages=[{"role": "user", "content": prompt}],
        temperature=0.3,
    )

    raw = response.choices[0].message.content.strip()

    lines = []
    for line in raw.splitlines():
        s = line.strip()
        if not s:
            continue
        s = re.sub(r"^(\d+\.|\*|-|•)\s*", "", s)
        if s:
            lines.append(s)

    if not lines:
        return original_bullets

    if len(lines) < bullet_count:
        lines.extend(original_bullets[len(lines):])
    elif len(lines) > bullet_count:
        lines = lines[:bullet_count]

    return lines


# ============================================================
#  Main resume update function
# ============================================================
def update_resume_experience(doc_bytes: BytesIO, job_description: str):
    """
    Load the uploaded DOCX, run ATS on the *current* text, then:

    - Rewrite Professional Experience bullets using missing keywords
    - Randomly replace some bullets with suggested bullets
    - Update SKILLS with categorized lines + Others
    - Recompute ATS

    Returns:
    - updated DOCX buffer
    - ATS score and analysis computed on the **updated** text
    - suggestions (extra bullets + skills) to display in UI
    """
    doc = docx.Document(doc_bytes)

    # 1) Get full text BEFORE rewrite
    original_text = read_full_text(doc)

    # 2) Compute ATS-style analysis on original resume
    score_before, kw_analysis_before, suggestions_before = compute_enterprise_ats(
        original_text, job_description
    )

    # 3) locate bullet paragraphs inside Professional Experience
    bullet_paras = get_experience_bullet_paragraphs(doc)
    if not bullet_paras:
        return None, score_before, kw_analysis_before, suggestions_before

    original_bullets = [p.text.strip() for p in bullet_paras]

    # 4) rewrite bullet contents using OpenAI
    new_bullets = rewrite_experience_bullets(
        original_bullets,
        job_description,
        kw_analysis_before,
    )

    # 4b) Inject suggested bullets at the END of the experience section
    #     (replace the last N bullets so they stay together and don't break roles)
    suggested_exp_bullets = suggestions_before.get("suggested_experience_bullets", []) or []
    if suggested_exp_bullets and new_bullets:
        max_replace = min(3, len(suggested_exp_bullets), len(new_bullets))
        if max_replace > 0:
            start_idx = len(new_bullets) - max_replace
            for offset, sug in enumerate(suggested_exp_bullets[:max_replace]):
                new_bullets[start_idx + offset] = sug

    # 5) apply back into the same paragraphs
    for para, new_text in zip(bullet_paras, new_bullets):
        replace_paragraph_text_preserving_format(para, new_text)

    # 6) update SKILLS section using suggestions + missing keywords
    update_skills_section(doc, suggestions_before, kw_analysis_before)

    # 7) save to BytesIO
    output_stream = BytesIO()
    doc.save(output_stream)
    output_stream.seek(0)

    # 8) recompute ATS on updated resume text
    updated_doc = docx.Document(output_stream)
    updated_text = read_full_text(updated_doc)

    score_after, kw_analysis_after, suggestions_after = compute_enterprise_ats(
        updated_text, job_description
    )

    return (
        output_stream,
        score_after,
        kw_analysis_after,
        suggestions_after,
    )


# ============================================================
#  Streamlit UI
# ============================================================
st.set_page_config(
    page_title="AI Resume Optimizer – Enterprise ATS + Experience Rewriter",
    layout="wide",
)

st.title("📄 AI Resume Optimizer – Enterprise ATS + Experience Rewriter")

st.markdown(
    """
This app:

- ✏️ Rewrites **Professional Experience bullet points** to match any job description  
- 🧾 Preserves your DOCX formatting, bullets, titles, projects, education  
- 🧠 Automatically **categorizes your Technical Skills** into clean labeled lines  
- ➕ Adds JD-aligned skills under **Others:** (e.g., Transformers, Feature Stores, Counterfactual Explainability, Personalized Ranking)  
- 📊 Computes an **enterprise-style ATS score** (core vs secondary vs nice-to-have keywords)  
- 💡 Uses LLM to generate strong bullets and missing skills, and injects some bullets **randomly** into your experience  

> Your resume should have a heading like **Professional Experience** and **Technical Skills / Skills**, and use real Word bullets.
"""
)

uploaded_file = st.file_uploader("Upload your resume (.docx)", type=["docx"])
job_description = st.text_area("Paste Job Description here", height=260)

col_btn1, col_btn2 = st.columns(2)

rewrite_clicked = False
ats_only_clicked = False

with col_btn1:
    if st.button("✨ Rewrite Experience + Update Skills + Run ATS"):
        rewrite_clicked = True

with col_btn2:
    if st.button("📊 Run ATS Analysis Only (No Rewriting)"):
        ats_only_clicked = True

if uploaded_file and job_description:
    if rewrite_clicked:
        with st.spinner("Rewriting Experience, updating Skills, and running ATS..."):
            result = update_resume_experience(uploaded_file, job_description)

        if result[0] is None:
            st.warning(
                "No bullet points found under 'Professional Experience'. "
                "Check the section heading or formatting."
            )
        else:
            optimized_buffer, score, kw_analysis, suggestions = result

            st.success("✅ Professional Experience bullets and Skills updated successfully.")

            st.download_button(
                "⬇ Download Optimized Resume",
                data=optimized_buffer,
                file_name="Resume_Optimized_Experience_Skills.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True,
            )

            # === ATS section ===
            st.markdown("---")
            st.subheader("📊 ATS Match Score (Enterprise-style)")

            st.metric("Estimated JD Match", f"{score}%")

            st.markdown("#### 🔑 JD Keyword Buckets")
            st.json(
                {
                    "core_keywords": kw_analysis["core_keywords"],
                    "secondary_keywords": kw_analysis["secondary_keywords"],
                    "nice_to_have": kw_analysis["nice_to_have"],
                }
            )

            st.markdown("#### ✅ Matched Keywords")
            core_present = kw_analysis["core_present"]
            sec_present = kw_analysis["secondary_present"]
            nice_present = kw_analysis["nice_present"]

            st.write(
                "**Core:** "
                + ", ".join(f"{k} (x{v})" for k, v in core_present.items())
                if core_present
                else "Core: none"
            )
            st.write(
                "**Secondary:** "
                + ", ".join(f"{k} (x{v})" for k, v in sec_present.items())
                if sec_present
                else "Secondary: none"
            )
            st.write(
                "**Nice-to-have:** "
                + ", ".join(f"{k} (x{v})" for k, v in nice_present.items())
                if nice_present
                else "Nice-to-have: none"
            )

            st.markdown("#### ❌ Missing / Weak Keywords (Good to cover)")
            st.write(
                "**Core:** "
                + ", ".join(kw_analysis["core_missing"])
                if kw_analysis["core_missing"]
                else "Core: none"
            )
            st.write(
                "**Secondary:** "
                + ", ".join(kw_analysis["secondary_missing"])
                if kw_analysis["secondary_missing"]
                else "Secondary: none"
            )
            st.write(
                "**Nice-to-have:** "
                + ", ".join(kw_analysis["nice_missing"])
                if kw_analysis["nice_missing"]
                else "Nice-to-have: none"
            )

            st.markdown("#### 💡 Suggested Bullets & Skills to Add (LLM Suggestions)")
            with st.expander("Suggested Experience Bullets"):
                for b in suggestions["suggested_experience_bullets"]:
                    st.markdown(f"- {b}")
            with st.expander("Suggested Skills to Add (also used under Others:)"):
                st.markdown(
                    ", ".join(suggestions["suggested_skills_to_add"])
                    if suggestions["suggested_skills_to_add"]
                    else "None"
                )

    elif ats_only_clicked:
        doc = docx.Document(uploaded_file)
        resume_text = read_full_text(doc)

        with st.spinner("Running ATS analysis on current resume..."):
            score, kw_analysis, suggestions = compute_enterprise_ats(
                resume_text, job_description
            )

        st.markdown("---")
        st.subheader("📊 ATS Match Score (Enterprise-style)")
        st.metric("Estimated JD Match", f"{score}%")

        st.markdown("#### 🔑 JD Keyword Buckets")
        st.json(
            {
                "core_keywords": kw_analysis["core_keywords"],
                "secondary_keywords": kw_analysis["secondary_keywords"],
                "nice_to_have": kw_analysis["nice_to_have"],
            }
        )

        st.markdown("#### ✅ Matched Keywords")
        core_present = kw_analysis["core_present"]
        sec_present = kw_analysis["secondary_present"]
        nice_present = kw_analysis["nice_present"]

        st.write(
            "**Core:** "
            + ", ".join(f"{k} (x{v})" for k, v in core_present.items())
            if core_present
            else "Core: none"
        )
        st.write(
            "**Secondary:** "
            + ", ".join(f"{k} (x{v})" for k, v in sec_present.items())
            if sec_present
            else "Secondary: none"
        )
        st.write(
            "**Nice-to-have:** "
            + ", ".join(f"{k} (x{v})" for k, v in nice_present.items())
            if nice_present
            else "Nice-to-have: none"
        )

        st.markdown("#### ❌ Missing / Weak Keywords (Good to cover)")
        st.write(
            "**Core:** "
            + ", ".join(kw_analysis["core_missing"])
            if kw_analysis["core_missing"]
            else "Core: none"
        )
        st.write(
            "**Secondary:** "
            + ", ".join(kw_analysis["secondary_missing"])
            if kw_analysis["secondary_missing"]
            else "Secondary: none"
        )
        st.write(
            "**Nice-to-have:** "
            + ", ".join(kw_analysis["nice_missing"])
            if kw_analysis["nice_missing"]
            else "Nice-to-have: none"
        )

        st.markdown("#### 💡 Suggested Bullets & Skills to Add (LLM Suggestions)")
        with st.expander("Suggested Experience Bullets"):
            for b in suggestions["suggested_experience_bullets"]:
                st.markdown(f"- {b}")
        with st.expander("Suggested Skills to Add"):
            st.markdown(
                ", ".join(suggestions["suggested_skills_to_add"])
                if suggestions["suggested_skills_to_add"]
                else "None"
            )